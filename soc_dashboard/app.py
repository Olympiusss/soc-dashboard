import streamlit as st
import pandas as pd
import datetime
import io
import json
import os
import base64
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import pyotp
import qrcode
import time
st.set_page_config(page_title="Attendance Checker", page_icon="🕒", layout="wide")
is_auth = st.session_state.get('authenticated', False)
page_id = "dashboard" if is_auth else "gateway"
st.markdown("""
<style>
/* --- Robust Curtain Animation --- */
@keyframes curtainLeft {
    0% { right: 50%; opacity: 1; z-index: 9999999; }
    99% { right: 100%; opacity: 0; z-index: 9999999; }
    100% { right: 100%; opacity: 0; z-index: -1; visibility: hidden; }
}
@keyframes curtainRight {
    0% { left: 50%; opacity: 1; z-index: 9999999; }
    99% { left: 100%; opacity: 0; z-index: 9999999; }
    100% { left: 100%; opacity: 0; z-index: -1; visibility: hidden; }
}
/* Dynamic curtain classes defined below */
/* Smoothen the entrance of the main content */
.block-container {
    animation: slideUpFade 1.0s ease forwards;
    animation-delay: 0.1s;
    opacity: 0;
}
@keyframes slideUpFade {
    0% { opacity: 0; transform: translateY(30px); }
    100% { opacity: 1; transform: translateY(0); }
}
/* File Uploader Hover Glow */
[data-testid="stFileUploader"] {
    border: 1px solid rgba(0, 198, 255, 0.2);
    border-radius: 12px;
    background: rgba(0, 114, 255, 0.02);
    padding: 10px;
    transition: all 0.3s ease;
}
[data-testid="stFileUploader"]:hover {
    border-color: #00c6ff;
    background: rgba(0, 198, 255, 0.05);
    box-shadow: 0 4px 15px rgba(0, 198, 255, 0.1);
}
/* Premium Card Hover Effects */
div[data-testid="stExpander"] {
    background-color: rgba(255, 255, 255, 0.02);
    border-radius: 8px;
    border: 1px solid rgba(255, 255, 255, 0.1);
    transition: all 0.3s ease;
}
div[data-testid="stExpander"]:hover {
    border-color: rgba(255, 255, 255, 0.3);
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.5);
}
/* Premium Primary Button */
.stButton>button {
    background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
    color: white !important;
    border-radius: 8px;
    border: none;
    font-weight: bold;
    transition: transform 0.2s, box-shadow 0.2s;
}
.stButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 15px rgba(42, 82, 152, 0.4);
}
/* Premium DataFrame border */
div[data-testid="stDataFrame"] {
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 4px 20px rgba(0,0,0,0.2) !important;
}
/* --- Output Pop-up Animations --- */
@keyframes outputPopUp {
    0%, 30% { transform: scale(0.9) translateY(40px); opacity: 0; }
    100% { transform: scale(1) translateY(0); opacity: 1; }
}
div[data-testid="stDataFrame"], .output-alert-box {
    animation: outputPopUp 1.2s cubic-bezier(0.175, 0.885, 0.32, 1.275) forwards;
}
/* Make headers look slightly crisper */
h1, h2, h3 {
    font-family: "Inter", sans-serif;
    letter-spacing: -0.5px;
}
</style>
""" + f"""
<style>
.curtain-panel-left-{page_id} {{
    position: fixed; top: 0; left: 0; right: 50%; bottom: 0;
    background-color: #0e1117; border-right: 2px solid rgba(0, 198, 255, 0.4);
    box-shadow: 5px 0 20px rgba(0,0,0,0.5);
    animation: curtainLeft 1.2s cubic-bezier(0.86, 0, 0.07, 1) forwards;
    animation-delay: 0.1s; z-index: 9999999; pointer-events: none;
}}
.curtain-panel-right-{page_id} {{
    position: fixed; top: 0; left: 50%; right: 0; bottom: 0;
    background-color: #0e1117; border-left: 2px solid rgba(0, 198, 255, 0.4);
    box-shadow: -5px 0 20px rgba(0,0,0,0.5);
    animation: curtainRight 1.2s cubic-bezier(0.86, 0, 0.07, 1) forwards;
    animation-delay: 0.1s; z-index: 9999999; pointer-events: none;
}}
</style>
<div class="curtain-panel-left-{page_id}"></div>
<div class="curtain-panel-right-{page_id}"></div>
""", unsafe_allow_html=True)
def get_header_html(logo_path, title_text, emoji_fallback, width=40, motto=""):
    title_style = "background: linear-gradient(135deg, #00c6ff 0%, #0072ff 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800; letter-spacing: -1px; line-height: 1.1;"
    
    motto_html = ""
    if motto:
        motto_html = f'<div style="margin-top: 8px; font-size: 1.15rem; background: linear-gradient(90deg, #89f7fe 0%, #66a6ff 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; border-left: 2px solid #89f7fe; padding-left: 12px; font-weight: 600; letter-spacing: 0.5px; display: inline-block;">✦ {motto}</div>'
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        # Flexbox layout with optional motto cleanly stacked under the title
        img_html = f'<img src="data:image/jpeg;base64,{b64}" width="{width}" style="margin-right: 18px; border-radius: 8px; box-shadow: 0 4px 15px rgba(0, 114, 255, 0.4); flex-shrink: 0; margin-top: 5px;">'
        return f'<div style="display: flex; align-items: flex-start; margin-bottom: 2.5rem;">{img_html}<div style="display: flex; flex-direction: column; justify-content: center;"><h1 style="margin: 0; padding: 0; {title_style}">{title_text}</h1>{motto_html}</div></div>'
    return f'<div style="margin-bottom: 2.5rem;"><h1 style="{title_style} margin: 0 0 10px 0;">{emoji_fallback} {title_text}</h1>{motto_html}</div>'
# --- Security Gateway (Google Authenticator) ---
if 'totp_secret' not in st.session_state:
    st.session_state['totp_secret'] = 'JBSWY3DPEBLW64TN'  # Can be moved to st.secrets later
if not st.session_state.get('authenticated', False):
    st.markdown(get_header_html(
        "sentrium_logo.jpg", 
        "Sentrium Security Gateway", 
        "🛡️", 
        width=55, 
        motto="We protect and strengthen the digital foundation of businesses"
    ), unsafe_allow_html=True)
    
    st.info("Please enter your 6-digit verification code from Google Authenticator.")
    
    totp = pyotp.TOTP(st.session_state['totp_secret'])
    
    with st.form("totp_form"):
        user_code = st.text_input("Enter 6-digit code", type="password")
        submit_button = st.form_submit_button("Verify")
        
        if submit_button:
            if totp.verify(user_code):
                st.session_state['authenticated'] = True
                st.rerun()
            else:
                st.error("Invalid code. Please try again.")
    st.stop()
st.markdown(get_header_html("sentrium_logo.jpg", "Sentrium Attendance Tracker", "🕒", width=60), unsafe_allow_html=True)
st.markdown(
    """<div style="background-color: rgba(0, 114, 255, 0.05); border-left: 4px solid #0072ff; padding: 15px 20px; border-radius: 0 8px 8px 0; margin-bottom: 2rem; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
    <span style="color: #e2e8f0; font-size: 1.05rem; line-height: 1.5;">Upload your access control logs alongside the <strong style="color: #00c6ff;">Staff Roster</strong> and/or the <strong style="color: #00c6ff;">SOC Schedule</strong> to generate a full attendance report — covering late arrivals and absentees for both departments.</span>
    </div>""", unsafe_allow_html=True
)
# ── Three-column upload layout ─────────────────────────────────────────────
up_col1, up_col2, up_col3 = st.columns(3)
with up_col1:
    uploaded_files = st.file_uploader(
        "Access Control Logs",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=True,
        help="Upload the raw badge-in logs from the door access system"
    )
with up_col2:
    roster_file = st.file_uploader(
        "Standard Staff Master List",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=False,
        help="Upload the monthly roster for standard staff (Mon–Fri). Enables Absentee tracking."
    )
with up_col3:
    soc_file = st.file_uploader(
        "SOC Team Schedule",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=False,
        help="Upload the SOC team's rotating shift schedule. Enables SOC-specific Late & Absent reporting."
    )
st.markdown("<br>", unsafe_allow_html=True)
# ── SOC Shift time mapping ──────────────────────────────────────────────────
SOC_SHIFT_TIMES = {
    'morning':   datetime.time(7, 0, 0),
    'afternoon': datetime.time(14, 0, 0),
    'night':     datetime.time(19, 0, 0),
}
# ── Helper: compact export dropdown ────────────────────────────────────────
def render_export_dropdown(dataframe, period_label, report_title, report_subtitle, col_widths_pdf, key_prefix):
    """Renders a compact dropdown with Excel/Word/PDF export options."""
    with st.popover("Export Report"):
        # Excel
        buf_xlsx = io.BytesIO()
        with pd.ExcelWriter(buf_xlsx, engine='openpyxl') as writer:
            dataframe.to_excel(writer, index=False, sheet_name=report_title[:31])
        st.download_button(
            label="Download as Excel (.xlsx)",
            data=buf_xlsx.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"xlsx_{key_prefix}_{period_label}"
        )
        # Word
        doc = Document()
        doc.add_heading(f'{report_title} — {period_label}', level=1)
        doc.add_paragraph(report_subtitle)
        tbl = doc.add_table(rows=1, cols=len(dataframe.columns), style='Light Grid Accent 1')
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for i, cn in enumerate(dataframe.columns):
            hdr[i].text = str(cn)
        for _, row in dataframe.iterrows():
            rc = tbl.add_row().cells
            for i, cn in enumerate(dataframe.columns):
                rc[i].text = str(row[cn])
        buf_docx = io.BytesIO()
        doc.save(buf_docx)
        st.download_button(
            label="Download as Word (.docx)",
            data=buf_docx.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{key_prefix}_{period_label}"
        )
        # PDF
        pdf = FPDF()
        pdf.add_page(orientation='L')
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, f'{report_title} - {period_label}', new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.set_font('Helvetica', '', 9)
        pdf.cell(0, 8, report_subtitle, new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.ln(5)
        pdf.set_font('Helvetica', 'B', 8)
        for i, h in enumerate(dataframe.columns):
            pdf.cell(col_widths_pdf[i] if i < len(col_widths_pdf) else 40, 8, str(h), border=1, align='C')
        pdf.ln()
        pdf.set_font('Helvetica', '', 7)
        for _, row in dataframe.iterrows():
            for i, cn in enumerate(dataframe.columns):
                w = col_widths_pdf[i] if i < len(col_widths_pdf) else 40
                pdf.cell(w, 7, str(row[cn]), border=1, align='C' if i in [0, 2] else 'L')
            pdf.ln()
        buf_pdf = io.BytesIO()
        pdf.output(buf_pdf)
        st.download_button(
            label="Download as PDF (.pdf)",
            data=buf_pdf.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.pdf",
            mime="application/pdf",
            key=f"pdf_{key_prefix}_{period_label}"
        )
# ── Helper: period filter (All / Daily / Weekly / Monthly) ──────────────────
def period_filter(df, date_col, parsed_time_col, radio_key):
    period = st.radio("View period:", ["All", "Daily", "Weekly", "Monthly"], horizontal=True, key=radio_key)
    if period == "Daily":
        df['_date_label'] = df[parsed_time_col].dt.date
        opts = sorted(df['_date_label'].unique(), reverse=True)
        sel = st.selectbox("Select Date", opts, key=f"{radio_key}_sel")
        return df[df['_date_label'] == sel].copy(), f"Day {sel}"
    elif period == "Weekly":
        df['_week'] = df[parsed_time_col].dt.isocalendar().year.astype(str) + '-W' + \
                      df[parsed_time_col].dt.isocalendar().week.astype(str).str.zfill(2)
        opts = sorted(df['_week'].unique(), reverse=True)
        sel = st.selectbox("Select Week", opts, key=f"{radio_key}_sel")
        return df[df['_week'] == sel].copy(), f"Week {sel}"
    elif period == "Monthly":
        df['_month'] = df[parsed_time_col].dt.to_period('M').astype(str)
        opts = sorted(df['_month'].unique(), reverse=True)
        sel = st.selectbox("Select Month", opts, key=f"{radio_key}_sel")
        return df[df['_month'] == sel].copy(), f"Month {sel}"
    return df.copy(), "All Time"
if uploaded_files:
    with st.spinner("Decrypting Access Logs & Compiling Report..."):
        time.sleep(1.2)
    try:
        # ── STEP 1: Load & normalise access logs ────────────────────────────
        STANDARD_COLS = {
            'time': 'Time', 'door name': 'Door Name',
            'event description': 'Event Description', 'personnel id': 'Personnel ID',
            'first name': 'First Name', 'last name': 'Last Name', 'door number': 'Door Number',
        }
        def normalize_columns(df_in):
            return df_in.rename(columns={
                c: STANDARD_COLS.get(str(c).strip().lower(), str(c).strip())
                for c in df_in.columns
            })
        def has_required_cols(cols):
            lc = [str(c).strip().lower() for c in cols]
            return all(t in lc for t in ['time', 'door name', 'event description'])
        all_dfs = []
        for uf in uploaded_files:
            if uf.name.endswith('.csv'):
                tmp = pd.read_csv(uf)
                tmp.columns = tmp.columns.astype(str).str.strip()
                all_dfs.append(normalize_columns(tmp))
            else:
                sheets = pd.read_excel(uf, sheet_name=None)
                for sname, d in sheets.items():
                    if d.empty:
                        continue
                    d.columns = d.columns.astype(str).str.strip()
                    if not has_required_cols(d.columns):
                        found = False
                        for i in range(min(20, len(d))):
                            rv = [str(v).strip() for v in d.iloc[i].values]
                            if has_required_cols(rv):
                                d.columns = rv
                                d = d.iloc[i+1:].reset_index(drop=True)
                                found = True
                                break
                        if not found:
                            continue
                    d = d.dropna(how='all').reset_index(drop=True)
                    d = normalize_columns(d)
                    if 'Time' in d.columns and 'Door Name' in d.columns:
                        all_dfs.append(d)
        if not all_dfs:
            st.error("No valid data found in the uploaded log files.")
            st.stop()
        df = pd.concat(all_dfs, ignore_index=True)
        # Validate
        for col in ['Event Description', 'Time', 'Door Name']:
            if col not in df.columns:
                st.error(f"Missing required column: **{col}**")
                st.stop()
        # Parse timestamps
        df['Parsed_Time'] = pd.to_datetime(df['Time'], errors='coerce', dayfirst=True)
        nat_mask = df['Parsed_Time'].isna()
        if nat_mask.any():
            df.loc[nat_mask, 'Parsed_Time'] = pd.to_datetime(df.loc[nat_mask, 'Time'], errors='coerce', dayfirst=False)
        df = df.dropna(subset=['Parsed_Time']).copy()
        df['Event Description'] = df['Event Description'].astype(str).str.strip()
        df['Door Name'] = df['Door Name'].astype(str).str.strip()
        # Fixed filters
        TARGET_DOORS  = ['Main entrance Ground Flr', 'Main entrance Upfloor']
        TARGET_EVENTS = ['Password', 'Normal Open']
        df_entry = df[
            df['Event Description'].str.lower().isin([e.lower() for e in TARGET_EVENTS]) &
            df['Door Name'].str.lower().isin([d.lower() for d in TARGET_DOORS])
        ].copy()
        df_entry['Date']      = df_entry['Parsed_Time'].dt.date
        df_entry['Time_Only'] = df_entry['Parsed_Time'].dt.time
        df_entry['Date_dt']   = pd.to_datetime(df_entry['Date'])
        # First badge-in per person per day
        df_entry = df_entry.sort_values('Parsed_Time')
        df_first = df_entry.drop_duplicates(subset=['Date', 'Personnel ID'], keep='first').copy()
        # ── STEP 2: Load rosters ────────────────────────────────────────────
        def load_simple_excel(file_obj):
            if file_obj.name.endswith('.csv'):
                return pd.read_csv(file_obj)
            return pd.read_excel(file_obj)
        df_roster = None
        soc_ids_from_roster = set()
        if roster_file:
            df_roster = load_simple_excel(roster_file)
            df_roster.columns = df_roster.columns.astype(str).str.strip()
            # Expect exactly: Personnel ID, Expected Days
            # Normalize known column names
            col_map = {}
            for c in df_roster.columns:
                cl = c.lower()
                if 'personnel' in cl or cl == 'id':
                    col_map[c] = 'Personnel ID'
                elif 'expected' in cl or 'days' in cl:
                    col_map[c] = 'Expected Days'
            df_roster = df_roster.rename(columns=col_map)
            if 'Personnel ID' not in df_roster.columns or 'Expected Days' not in df_roster.columns:
                st.warning("Master List must contain 'Personnel ID' and 'Expected Days' columns. Absentee tracking disabled.")
                df_roster = None
            else:
                df_roster['Personnel ID'] = df_roster['Personnel ID'].astype(str).str.strip()
                df_roster['Expected Days'] = df_roster['Expected Days'].astype(str).str.strip()
                # Identify SOC staff (Expected Days = "SOC")
                soc_mask = df_roster['Expected Days'].str.upper() == 'SOC'
                soc_ids_from_roster = set(df_roster.loc[soc_mask, 'Personnel ID'].unique())
                # Keep only standard staff for absentee logic
                df_std_roster = df_roster[~soc_mask].copy()
                df_std_roster['Days_Per_Week'] = pd.to_numeric(df_std_roster['Expected Days'], errors='coerce')
                df_std_roster = df_std_roster.dropna(subset=['Days_Per_Week']).copy()
                df_std_roster['Days_Per_Week'] = df_std_roster['Days_Per_Week'].astype(int)
        df_soc = None
        if soc_file:
            df_soc = load_simple_excel(soc_file)
            df_soc.columns = df_soc.columns.astype(str).str.strip()
            df_soc = normalize_columns(df_soc)
            # Expect: Personnel ID, Date, Shift (+ optional First Name, Last Name)
            for needed in ['Personnel ID', 'Date', 'Shift']:
                if needed not in df_soc.columns:
                    st.warning(f"SOC Schedule must contain a '{needed}' column. SOC tab disabled.")
                    df_soc = None
                    break
            if df_soc is not None:
                df_soc['Date'] = pd.to_datetime(df_soc['Date'], errors='coerce', dayfirst=True).dt.date
                df_soc = df_soc.dropna(subset=['Date']).copy()
                df_soc['Shift_Lower'] = df_soc['Shift'].astype(str).str.strip().str.lower()
                df_soc['Shift_Start'] = df_soc['Shift_Lower'].map(SOC_SHIFT_TIMES)
                df_soc['Personnel ID'] = df_soc['Personnel ID'].astype(str).str.strip()
                fn_cols = [c for c in ['First Name', 'Last Name'] if c in df_soc.columns]
                df_soc['Full Name'] = df_soc[fn_cols].fillna('').astype(str).apply(
                    lambda r: ' '.join(r).strip(), axis=1) if fn_cols else ''
        # ── STEP 3: Tabs ────────────────────────────────────────────────────
        st.divider()
        tab_std, tab_soc = st.tabs(["Standard Staff", "SOC Team"])
        # ════════════════════════════════════════════════════════════════════
        # TAB 1 — STANDARD STAFF
        # ════════════════════════════════════════════════════════════════════
        with tab_std:
            CUTOFF = datetime.time(8, 30, 0)
            WORKDAYS = {0, 1, 2, 3, 4}  # Mon–Fri
            # Combine SOC IDs from master list + SOC schedule
            soc_ids = soc_ids_from_roster.copy()
            if df_soc is not None:
                soc_ids.update(df_soc['Personnel ID'].astype(str).str.strip().unique())
            df_std_first = df_first[~df_first['Personnel ID'].astype(str).str.strip().isin(soc_ids)].copy()
            # ── Late Arrivals ────────────────────────────────────────────────
            st.subheader("Late Arrivals — Standard Staff")
            # Filter Mon-Fri only
            df_std_workday = df_std_first[df_std_first['Date_dt'].dt.dayofweek.isin(WORKDAYS)].copy()
            df_late = df_std_workday[df_std_workday['Time_Only'] >= CUTOFF].copy()
            # Build name columns
            name_cols = [c for c in ['First Name', 'Last Name'] if c in df_late.columns]
            df_late['Full Name'] = df_late[name_cols].fillna('').astype(str).apply(
                lambda r: ' '.join(r).strip(), axis=1) if name_cols else ''
            df_late['Date_Str'] = df_late['Date'].apply(lambda d: d.strftime('%Y-%m-%d'))
            df_late['Time_Str'] = df_late['Time_Only'].apply(lambda t: t.strftime('%H:%M:%S'))
            if df_late.empty:
                st.markdown(
                    '<div class="output-alert-box" style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                    '<h3 style="margin:0;color:#30d158;">Perfect Attendance</h3>'
                    '<p style="margin:5px 0 0;color:#e2e8f0;">No standard staff checked in after 8:30 AM.</p></div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    f'<div class="output-alert-box" style="background:rgba(255,69,58,0.1);border-left:5px solid #ff453a;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                    f'<h3 style="margin:0;color:#ff6b6b;">{len(df_late)} Late Check-in Record(s)</h3>'
                    f'<p style="margin:5px 0 0;color:#e2e8f0;">Standard staff who arrived at or after <strong>8:30 AM</strong> on a working day.</p></div>',
                    unsafe_allow_html=True
                )
                # Period filter row + export dropdown on the right
                pf_col, ex_col = st.columns([4, 1])
                with pf_col:
                    df_late_f, period_label_late = period_filter(df_late.copy(), 'Date', 'Parsed_Time', 'std_late_period')
                if df_late_f.empty:
                    st.info("No late arrivals for the selected period.")
                else:
                    # Build the aggregated table with Expected/Week
                    # Add ISO week to compute weekly default count
                    df_late_f['ISO_Week'] = df_late_f['Parsed_Time'].dt.isocalendar().year.astype(str) + '-W' + \
                        df_late_f['Parsed_Time'].dt.isocalendar().week.astype(str).str.zfill(2)
                    late_agg = (
                        df_late_f.sort_values(['Date', 'Time_Only'])
                        .groupby(['Personnel ID', 'Full Name'])
                        .agg(
                            Default_Count=('Date', 'nunique'),
                            Weeks_Defaulted=('ISO_Week', 'nunique'),
                            Dates_Defaulted=('Date_Str', lambda x: ', '.join(x)),
                            Times_Defaulted=('Time_Str', lambda x: ', '.join(x)),
                        )
                        .reset_index()
                        .sort_values('Default_Count', ascending=False)
                    )
                    # Attach Expected Days/Week from roster if available
                    if df_roster is not None:
                        late_agg = late_agg.merge(
                            df_std_roster[['Personnel ID', 'Days_Per_Week']],
                            on='Personnel ID', how='left'
                        )
                        late_agg['Days_Per_Week'] = late_agg['Days_Per_Week'].fillna('–')
                        late_agg = late_agg[['Personnel ID', 'Full Name', 'Days_Per_Week',
                                            'Default_Count', 'Weeks_Defaulted', 'Dates_Defaulted', 'Times_Defaulted']]
                        late_agg.columns = ['Personnel ID', 'Full Name', 'Expected/Week',
                                            'Days Late', 'Weeks Defaulted', 'Dates Late', 'Times In']
                    else:
                        late_agg.columns = ['Personnel ID', 'Full Name', 'Days Late', 'Weeks Defaulted', 'Dates Late', 'Times In']
                    st.dataframe(late_agg, hide_index=True, use_container_width=True)
                    with ex_col:
                        render_export_dropdown(
                            late_agg, period_label_late,
                            "Standard Staff Late Arrivals",
                            "Personnel who checked in at or after 08:30 AM on working days.",
                            [20, 40, 15, 15, 15, 80, 80], "std_late"
                        )
            # ── Absentees (Weekly Compliance) ─────────────────────────────────
            st.divider()
            st.subheader("Weekly Attendance Compliance — Standard Staff")
            if df_roster is None:
                st.info("Upload the Standard Staff Master List to enable weekly compliance tracking.")
            else:
                # Build name lookup from access logs
                name_lookup = df_first[['Personnel ID', 'First Name', 'Last Name']].copy() if \
                    'First Name' in df_first.columns else None
                if name_lookup is not None:
                    name_lookup['Personnel ID'] = name_lookup['Personnel ID'].astype(str).str.strip()
                    nc = [c for c in ['First Name', 'Last Name'] if c in name_lookup.columns]
                    name_lookup['Full Name'] = name_lookup[nc].fillna('').astype(str).apply(
                        lambda r: ' '.join(r).strip(), axis=1)
                    name_lookup = name_lookup.drop_duplicates(subset='Personnel ID')[['Personnel ID', 'Full Name']]
                # Get actual attendance per person per ISO week
                df_std_attendance = df_std_first.copy()
                df_std_attendance['Personnel ID'] = df_std_attendance['Personnel ID'].astype(str).str.strip()
                df_std_attendance['ISO_Year'] = df_std_attendance['Parsed_Time'].dt.isocalendar().year.astype(int)
                df_std_attendance['ISO_Week'] = df_std_attendance['Parsed_Time'].dt.isocalendar().week.astype(int)
                df_std_attendance['Week_Label'] = df_std_attendance['ISO_Year'].astype(str) + '-W' + \
                    df_std_attendance['ISO_Week'].astype(str).str.zfill(2)
                # Count unique days per person per week
                actual_weekly = (
                    df_std_attendance
                    .groupby(['Personnel ID', 'Week_Label'])
                    .agg(Actual_Days=('Date', 'nunique'))
                    .reset_index()
                )
                # Build all expected weeks from the log date range
                min_dt = df_std_attendance['Parsed_Time'].min()
                max_dt = df_std_attendance['Parsed_Time'].max()
                all_weeks = pd.date_range(min_dt, max_dt, freq='W-MON')
                week_labels = []
                for w in all_weeks:
                    iso = w.isocalendar()
                    week_labels.append(f"{iso[0]}-W{str(iso[1]).zfill(2)}")
                if not week_labels:
                    # Single partial week
                    iso = min_dt.isocalendar()
                    week_labels = [f"{iso[0]}-W{str(iso[1]).zfill(2)}"]
                # Cross-product: every standard staff person × every week
                expected_weeks = pd.MultiIndex.from_product(
                    [df_std_roster['Personnel ID'].unique(), week_labels],
                    names=['Personnel ID', 'Week_Label']
                ).to_frame(index=False)
                # Merge expected days per week from roster
                expected_weeks = expected_weeks.merge(
                    df_std_roster[['Personnel ID', 'Days_Per_Week']],
                    on='Personnel ID', how='left'
                )
                # Merge actual attendance
                compliance = expected_weeks.merge(actual_weekly, on=['Personnel ID', 'Week_Label'], how='left')
                compliance['Actual_Days'] = compliance['Actual_Days'].fillna(0).astype(int)
                compliance['Deficit'] = compliance['Days_Per_Week'] - compliance['Actual_Days']
                # Only flag non-compliant weeks (deficit > 0)
                non_compliant = compliance[compliance['Deficit'] > 0].copy()
                # Attach names
                if name_lookup is not None:
                    non_compliant = non_compliant.merge(name_lookup, on='Personnel ID', how='left')
                    non_compliant['Full Name'] = non_compliant['Full Name'].fillna('')
                else:
                    non_compliant['Full Name'] = ''
                if non_compliant.empty:
                    st.markdown(
                        '<div style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;">'
                        '<h3 style="margin:0;color:#30d158;">Full Compliance</h3>'
                        '<p style="margin:5px 0 0;color:#e2e8f0;">All standard staff met their required weekly attendance.</p></div>',
                        unsafe_allow_html=True
                    )
                else:
                    total_nc_staff = non_compliant['Personnel ID'].nunique()
                    st.markdown(
                        f'<div style="background:rgba(255,149,0,0.1);border-left:5px solid #ff9500;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                        f'<h3 style="margin:0;color:#ff9500;">{total_nc_staff} staff member(s) had non-compliant weeks</h3>'
                        f'<p style="margin:5px 0 0;color:#e2e8f0;">Based on required days per week vs. actual badge-in records.</p></div>',
                        unsafe_allow_html=True
                    )
                    # Detailed weekly breakdown
                    st.markdown("**Weekly Detail** — each row is a week where the employee fell short.")
                    detail_df = non_compliant[['Personnel ID', 'Full Name', 'Week_Label', 'Days_Per_Week', 'Actual_Days', 'Deficit']].copy()
                    detail_df.columns = ['Personnel ID', 'Full Name', 'Week', 'Expected', 'Actual', 'Deficit']
                    detail_df = detail_df.sort_values(['Personnel ID', 'Week'])
                    st.dataframe(detail_df, hide_index=True, use_container_width=True)
                    # Aggregated summary
                    st.markdown("**Summary** — total non-compliant weeks and deficit days per employee.")
                    summary_agg = (
                        non_compliant
                        .groupby(['Personnel ID', 'Full Name'])
                        .agg(
                            Weeks_Non_Compliant=('Week_Label', 'nunique'),
                            Total_Deficit=('Deficit', 'sum'),
                            Weeks_List=('Week_Label', lambda x: ', '.join(sorted(x))),
                        )
                        .reset_index()
                        .sort_values('Total_Deficit', ascending=False)
                    )
                    summary_agg.columns = ['Personnel ID', 'Full Name', 'Weeks Non-Compliant', 'Total Deficit Days', 'Weeks']
                    st.dataframe(summary_agg, hide_index=True, use_container_width=True)
                    render_export_dropdown(
                        summary_agg, "All Time",
                        "Weekly Attendance Compliance",
                        "Standard staff who did not meet their required days per week.",
                        [20, 45, 25, 25, 80], "std_compliance"
                    )
        # ════════════════════════════════════════════════════════════════════
        # TAB 2 — SOC TEAM
        # ════════════════════════════════════════════════════════════════════
        with tab_soc:
            if df_soc is None:
                st.info("Upload the SOC Team Schedule file to enable the SOC Department report.")
                st.markdown("""
                **Expected columns in the SOC Schedule file:**
                | Personnel ID | First Name | Last Name | Date | Shift |
                |---|---|---|---|---|
                | 201 | Jane | Smith | 14/04/2026 | Morning |
                | 201 | Jane | Smith | 15/04/2026 | Night |
                | 202 | James | Okoro | 14/04/2026 | Afternoon |
                **Shift options:** `Morning` (07:00) · `Afternoon` (14:00) · `Night` (19:00)
                > Days off should simply be omitted — do not add a row for days off.
                """)
            else:
                # Cross-reference SOC schedule with access logs
                df_soc_pid = df_soc['Personnel ID'].unique()
                df_soc_logs = df_first[df_first['Personnel ID'].astype(str).str.strip().isin(df_soc_pid)].copy()
                df_soc_logs['Personnel ID'] = df_soc_logs['Personnel ID'].astype(str).str.strip()
                # Merge schedule with actual logs
                merged_soc = df_soc.merge(
                    df_soc_logs[['Personnel ID', 'Date', 'Time_Only', 'Parsed_Time']],
                    on=['Personnel ID', 'Date'],
                    how='left'
                )
                # ── SOC Late Arrivals ───────────────────────────────────────
                st.subheader("Late Arrivals — SOC Team")
                soc_present = merged_soc.dropna(subset=['Time_Only']).copy()
                soc_present['Shift_Start_dt'] = soc_present.apply(
                    lambda r: datetime.datetime.combine(r['Date'], r['Shift_Start'])
                    if pd.notna(r['Shift_Start']) else None, axis=1
                )
                soc_present['Is_Late'] = soc_present.apply(
                    lambda r: r['Time_Only'] > r['Shift_Start'] if r['Shift_Start'] else False, axis=1
                )
                soc_present['Minutes_Late'] = soc_present.apply(
                    lambda r: int(
                        (datetime.datetime.combine(r['Date'], r['Time_Only']) -
                         datetime.datetime.combine(r['Date'], r['Shift_Start'])).total_seconds() // 60
                    ) if r['Is_Late'] and r['Shift_Start'] else 0, axis=1
                )
                soc_late = soc_present[soc_present['Is_Late']].copy()
                soc_late['Date_Str'] = soc_late['Date'].apply(lambda d: d.strftime('%Y-%m-%d'))
                soc_late['Shift_Start_Str'] = soc_late['Shift_Start'].apply(lambda t: t.strftime('%H:%M') if pd.notna(t) else '')
                soc_late['Time_In_Str'] = soc_late['Time_Only'].apply(lambda t: t.strftime('%H:%M:%S'))
                if soc_late.empty:
                    st.markdown(
                        '<div class="output-alert-box" style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                        '<h3 style="margin:0;color:#30d158;">✅ Perfect Punctuality!</h3>'
                        '<p style="margin:5px 0 0;color:#e2e8f0;">No SOC team members were late for their scheduled shifts.</p></div>',
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f'<div class="output-alert-box" style="background:rgba(255,69,58,0.1);border-left:5px solid #ff453a;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                        f'<h3 style="margin:0;color:#ff6b6b;">⚠️ {len(soc_late)} Late Shift Arrival(s)</h3>'
                        f'<p style="margin:5px 0 0;color:#e2e8f0;">SOC members who arrived after their scheduled shift start time.</p></div>',
                        unsafe_allow_html=True
                    )
                    soc_late_f, pl_soc_late = period_filter(soc_late.copy(), 'Date', 'Parsed_Time', 'soc_late_period')
                    if soc_late_f.empty:
                        st.info("No SOC late arrivals for the selected period.")
                    else:
                        soc_late_agg = (
                            soc_late_f.sort_values('Date')
                            .groupby(['Personnel ID', 'Full Name'])
                            .agg(
                                Late_Count=('Date', 'nunique'),
                                Dates_Late=('Date_Str', lambda x: ', '.join(x)),
                                Shifts=('Shift', lambda x: ', '.join(x)),
                                Scheduled_Start=('Shift_Start_Str', lambda x: ', '.join(x)),
                                Times_In=('Time_In_Str', lambda x: ', '.join(x)),
                                Total_Mins_Late=('Minutes_Late', 'sum'),
                            )
                            .reset_index()
                            .sort_values('Late_Count', ascending=False)
                        )
                        soc_late_agg.columns = [
                            'Personnel ID', 'Full Name', 'Late Count',
                            'Dates Late', 'Shifts', 'Scheduled Start', 'Time In', 'Total Mins Late'
                        ]
                        st.dataframe(soc_late_agg, hide_index=True, use_container_width=True)
                        render_export_dropdown(
                            soc_late_agg, pl_soc_late,
                            "SOC Team Late Arrivals",
                            "SOC personnel who arrived after their scheduled shift start time.",
                            [18, 40, 15, 50, 30, 25, 30, 20], "soc_late"
                        )
                # ── SOC Absentees ───────────────────────────────────────────
                st.divider()
                st.subheader("Absentees — SOC Team")
                soc_absent = merged_soc[merged_soc['Time_Only'].isna()].copy()
                soc_absent['Date_Str'] = soc_absent['Date'].apply(lambda d: d.strftime('%Y-%m-%d'))
                soc_absent['Date_dt_col'] = pd.to_datetime(soc_absent['Date'])
                if soc_absent.empty:
                    st.markdown(
                        '<div style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;">'
                        '<h3 style="margin:0;color:#30d158;">✅ No SOC Absentees!</h3>'
                        '<p style="margin:5px 0 0;color:#e2e8f0;">All SOC members were present for every scheduled shift.</p></div>',
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f'<div style="background:rgba(255,149,0,0.1);border-left:5px solid #ff9500;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                        f'<h3 style="margin:0;color:#ff9500;">🚫 {soc_absent["Personnel ID"].nunique()} SOC member(s) missed scheduled shifts</h3>'
                        f'<p style="margin:5px 0 0;color:#e2e8f0;">No badge-in recorded on their scheduled shift day.</p></div>',
                        unsafe_allow_html=True
                    )
                    soc_absent_f, pl_soc_abs = period_filter(soc_absent.copy(), 'Date', 'Date_dt_col', 'soc_abs_period')
                    if soc_absent_f.empty:
                        st.info("No SOC absences for the selected period.")
                    else:
                        soc_abs_agg = (
                            soc_absent_f.sort_values('Date')
                            .groupby(['Personnel ID', 'Full Name'])
                            .agg(
                                Absent_Count=('Date', 'nunique'),
                                Dates_Absent=('Date_Str', lambda x: ', '.join(sorted(x))),
                                Shifts_Missed=('Shift', lambda x: ', '.join(x)),
                            )
                            .reset_index()
                            .sort_values('Absent_Count', ascending=False)
                        )
                        soc_abs_agg.columns = [
                            'Personnel ID', 'Full Name', 'Shifts Missed', 'Dates Absent', 'Shifts'
                        ]
                        st.dataframe(soc_abs_agg, hide_index=True, use_container_width=True)
                        render_export_dropdown(
                            soc_abs_agg, pl_soc_abs,
                            "SOC Team Absentees",
                            "SOC personnel with no badge-in on their scheduled shift days.",
                            [18, 40, 20, 80, 50], "soc_absent"
                        )
    except Exception as e:
        st.error(f"An error occurred while processing: {e}")
